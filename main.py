from __future__ import annotations
import asyncio
from typing import List, Optional
from datetime import datetime
import os
from pathlib import Path
import json
import shutil

from fastapi import FastAPI, Depends, HTTPException, status, Request, Form, UploadFile, File
from fastapi.responses import HTMLResponse, RedirectResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from pydantic import BaseModel
from sqlalchemy import Column, Integer, String, JSON, DateTime, func, select
from sqlalchemy.ext.asyncio import create_async_engine, AsyncSession, async_sessionmaker
from sqlalchemy.orm import declarative_base

# NOTE: This file expects these third-party packages installed:
# pip install fastapi uvicorn sqlalchemy aiosqlite python-docx pdfminer.six
# (pdf text extraction uses pdfminer.high_level.extract_text)

# ----------------------------
# Paths and static/templates setup
# ----------------------------
BASE_DIR = Path(__file__).resolve().parent
TEMPLATES_DIR = BASE_DIR / "templates"
STATIC_DIR = BASE_DIR / "static"
UPLOADS_DIR = BASE_DIR / "uploads"
TEMPLATES_DIR.mkdir(exist_ok=True)
STATIC_DIR.mkdir(exist_ok=True)
UPLOADS_DIR.mkdir(exist_ok=True)

# index.html — note: removed free-text input; upload only
INDEX_HTML = """<!doctype html>
<html>
<head>
  <meta charset="utf-8">
  <title>Documents — Список</title>
  <link rel="stylesheet" href="/static/styles.css">
</head>
<body>
  <div class="container">
    <h1>Список документов</h1>

    {% if documents %}
      <table class="docs">
        <thead><tr><th>ID</th><th>file_name</th><th>name</th><th>created_at</th><th>actions</th></tr></thead>
        <tbody>
        {% for d in documents %}
          <tr>
            <td>{{ d.id }}</td>
            <td><a href="/documents/{{ d.id }}">{{ d.file_name }}</a></td>
            <td>{{ d.name or "-" }}</td>
            <td>{{ d.created_at or "-" }}</td>
            <td>
              <form action="/documents/{{ d.id }}/delete" method="post" style="display:inline">
                <button type="submit" onclick="return confirm('Удалить документ {{ d.file_name }}?')">Удалить</button>
              </form>
            </td>
          </tr>
        {% endfor %}
        </tbody>
      </table>
    {% else %}
      <p>Документов пока нет.</p>
    {% endif %}

    <hr/>
    <h2>Загрузить документ (PDF или DOCX)</h2>
    <form action="/documents/upload" method="post" enctype="multipart/form-data">
      <label>Файл:<br/>
        <input type="file" name="file" accept=".pdf,.docx" required>
      </label>
      <br/>
      <label>Имя (опционально):<br/>
        <input type="text" name="name" style="width: 100%">
      </label>
      <br/>
      <button type="submit">Загрузить</button>
    </form>

    <footer>
      <p>Пример: FastAPI + Jinja2. Файл обрабатывается и извлекается текст автоматически.</p>
    </footer>
  </div>
</body>
</html>
"""

# document.html and error.html and static assets copied from original, minor safe adjustments
DOCUMENT_HTML = """<!doctype html>
<html>
<head>
  <meta charset="utf-8">
  <title>Документ {{ doc.file_name }}</title>
  <link rel="stylesheet" href="/static/styles.css">
  <script src="/static/toggle.js"></script>
</head>
<body>
  <div class="container">
    <a href="/">← Назад к списку</a>
    <h1>{{ doc.file_name }} {% if doc.name %} — {{ doc.name }}{% endif %}</h1>
    <p><small>id: {{ doc.id }} — создан: {{ doc.created_at or "-" }}</small></p>

    <h2>Оригинальный текст</h2>
    <pre class="original">{{ doc.original_text }}</pre>

    <h2>Сводки (summary)</h2>

    <div class="summary-block">
      <button onclick="toggle('llm_text')" class="toggle-btn">LLM Text Summary (RU/EN)</button>
      <div id="llm_text" class="collapsible">
        <h3>RU</h3>
        <p>{{ doc.summary.llm_text_summary.ru }}</p>
        <h3>EN</h3>
        <p>{{ doc.summary.llm_text_summary.en }}</p>
      </div>
    </div>

    <div class="summary-block">
      <button onclick="toggle('extr_text')" class="toggle-btn">Extraction Text Summary (RU/EN)</button>
      <div id="extr_text" class="collapsible">
        <h3>RU</h3>
        <p>{{ doc.summary.extraction_text_summary.ru }}</p>
        <h3>EN</h3>
        <p>{{ doc.summary.extraction_text_summary.en }}</p>
      </div>
    </div>

    <div class="summary-block">
      <button onclick="toggle('llm_kw')" class="toggle-btn">LLM Keyword Tree</button>
      <div id="llm_kw" class="collapsible">
        {% macro render_node(node) -%}
          <li>{{ node.keyword }}
            {% if node.children %}
              <ul>
                {% for c in node.children %}
                  {{ render_node(c) }}
                {% endfor %}
              </ul>
            {% endif %}
          </li>
        {%- endmacro %}
        <ul class="kw-tree">
          {{ render_node(doc.summary.llm_keyword_summary.ru) }}
        </ul>
      </div>
    </div>

    <div class="summary-block">
      <button onclick="toggle('extr_kw')" class="toggle-btn">Extraction Keyword Tree</button>
      <div id="extr_kw" class="collapsible">
        {% macro render_node2(node) -%}
          <li>{{ node.keyword }}
            {% if node.children %}
              <ul>
                {% for c in node.children %}
                  {{ render_node2(c) }}
                {% endfor %}
              </ul>
            {% endif %}
          </li>
        {%- endmacro %}
        <ul class="kw-tree">
          {{ render_node2(doc.summary.extraction_keyword_summary.ru) }}
        </ul>
      </div>
    </div>

    <form action="/documents/{{ doc.id }}/delete" method="post" style="margin-top:1em;">
      <button type="submit" onclick="return confirm('Удалить документ?')">Удалить документ</button>
    </form>

  </div>
</body>
</html>
"""

ERROR_HTML = """<!doctype html>
<html>
<head>
  <meta charset="utf-8">
  <title>Ошибка</title>
  <link rel="stylesheet" href="/static/styles.css">
</head>
<body>
  <div class="container">
    <h1>Ошибка</h1>
    <p>{{ message }}</p>
    <p><a href="/">Вернуться на главную</a></p>
  </div>
</body>
</html>
"""

STYLES_CSS = """
body { font-family: Arial, sans-serif; background:#f8f9fb; color:#222; }
.container { max-width:900px; margin:24px auto; background:white; padding:18px; box-shadow:0 2px 6px rgba(0,0,0,0.08); border-radius:8px; }
.docs { width:100%; border-collapse: collapse; margin-bottom:12px; }
.docs th, .docs td { border:1px solid #ddd; padding:8px; text-align:left; }
.original { white-space:pre-wrap; background:#f1f1f1; padding:12px; border-radius:6px; }
.summary-block { margin-top:12px; }
.toggle-btn { background:#007bff; color:white; padding:6px 10px; border-radius:6px; border:none; cursor:pointer; }
.collapsible { display:none; margin-top:8px; padding:10px; border-left:3px solid #ddd; background:#fcfcff; border-radius:6px; }
.kw-tree { list-style-type: disc; margin-left: 18px; }
footer { margin-top:18px; color:#666; font-size:0.9em; }
"""

TOGGLE_JS = """
function toggle(id) {
  const el = document.getElementById(id);
  if (!el) return;
  if (el.style.display === 'block') el.style.display = 'none';
  else el.style.display = 'block';
}
"""

# Write files if not present (templates + static)
def ensure_static_and_templates():
    if not (TEMPLATES_DIR / "index.html").exists():
        (TEMPLATES_DIR / "index.html").write_text(INDEX_HTML, encoding="utf-8")
    if not (TEMPLATES_DIR / "document.html").exists():
        (TEMPLATES_DIR / "document.html").write_text(DOCUMENT_HTML, encoding="utf-8")
    if not (TEMPLATES_DIR / "error.html").exists():
        (TEMPLATES_DIR / "error.html").write_text(ERROR_HTML, encoding="utf-8")
    if not (STATIC_DIR / "styles.css").exists():
        (STATIC_DIR / "styles.css").write_text(STYLES_CSS, encoding="utf-8")
    if not (STATIC_DIR / "toggle.js").exists():
        (STATIC_DIR / "toggle.js").write_text(TOGGLE_JS, encoding="utf-8")

ensure_static_and_templates()

# ----------------------------
# DTOs & forward refs
# ----------------------------
class KeywordNode(BaseModel):
    keyword: str
    children: List["KeywordNode"] = []

class TextSummary(BaseModel):
    ru: str
    en: str

class KeywordTreeSummary(BaseModel):
    ru: KeywordNode
    en: KeywordNode

class SummaryResult(BaseModel):
    llm_text_summary: TextSummary
    llm_keyword_summary: KeywordTreeSummary
    extraction_text_summary: TextSummary
    extraction_keyword_summary: KeywordTreeSummary

class TextDocumentDTO(BaseModel):
    id: int
    file_name: str
    name: Optional[str]
    original_text: str
    summary: SummaryResult

class DocumentInfoDTO(BaseModel):
    id: int
    file_name: str
    name: Optional[str]
    created_at: Optional[datetime]

KeywordNode.update_forward_refs()
KeywordTreeSummary.update_forward_refs()
SummaryResult.update_forward_refs()

# ----------------------------
# SQLAlchemy ORM
# ----------------------------
Base = declarative_base()

class TextDocument(Base):
    __tablename__ = "text_documents"
    id = Column(Integer, primary_key=True, autoincrement=True)
    file_name = Column(String, nullable=False, unique=True)
    name = Column(String, nullable=True)
    original_text = Column(String, nullable=False)
    summary_json = Column(JSON, nullable=False)
    created_at = Column(DateTime(timezone=True), server_default=func.now())

# ----------------------------
# Async repository
# ----------------------------
class TextRepositoryAsync:
    def __init__(self, db_url: str = "sqlite+aiosqlite:///texts_async.db") -> None:
        self.engine = create_async_engine(db_url, echo=False, future=True)
        self.async_session = async_sessionmaker(self.engine, expire_on_commit=False, class_=AsyncSession)

    async def init_models(self) -> None:
        async with self.engine.begin() as conn:
            await conn.run_sync(Base.metadata.create_all)

    async def add_document(
        self,
        original_text: str,
        summary_result: SummaryResult,
        file_name: str,
        name: Optional[str] = None
    ) -> int:
        async with self.async_session() as session:
            async with session.begin():
                doc = TextDocument(
                    original_text=original_text,
                    summary_json=summary_result.dict(),
                    file_name=file_name,
                    name=name
                )
                session.add(doc)
            await session.refresh(doc)
            return doc.id

    async def get_document(self, doc_id: int) -> Optional[TextDocumentDTO]:
        async with self.async_session() as session:
            doc = await session.get(TextDocument, doc_id)
            if doc is None:
                return None
            summary_obj = SummaryResult(**doc.summary_json)
            return TextDocumentDTO(
                id=doc.id,
                file_name=doc.file_name,
                name=doc.name,
                original_text=doc.original_text,
                summary=summary_obj
            )

    async def list_documents(self) -> List[TextDocumentDTO]:
        async with self.async_session() as session:
            result = await session.execute(select(TextDocument))
            docs = result.scalars().all()
            return [
                TextDocumentDTO(
                    id=d.id,
                    file_name=d.file_name,
                    name=d.name,
                    original_text=d.original_text,
                    summary=SummaryResult(**d.summary_json)
                )
                for d in docs
            ]

    async def list_document_info(self) -> List[DocumentInfoDTO]:
        async with self.async_session() as session:
            result = await session.execute(select(TextDocument))
            docs = result.scalars().all()
            return [
                DocumentInfoDTO(
                    id=d.id,
                    file_name=d.file_name,
                    name=d.name,
                    created_at=d.created_at
                ) for d in docs
            ]

    async def delete_document(self, doc_id: int) -> bool:
        async with self.async_session() as session:
            async with session.begin():
                doc = await session.get(TextDocument, doc_id)
                if doc is None:
                    return False
                await session.delete(doc)
        return True

# ----------------------------
# Sub-services (stubs)
# ----------------------------
class LLMTextSummaryService:
    def __init__(self) -> None:
        pass

    async def generate(self, text: str) -> TextSummary:
        await asyncio.sleep(0.1)
        return TextSummary(ru="LLM: краткое резюме (RU)", en="LLM: short summary (EN)")

class ExtractionTextSummaryService:
    def __init__(self) -> None:
        pass

    async def generate(self, text: str) -> TextSummary:
        await asyncio.sleep(0.05)
        return TextSummary(ru="Extraction: краткое резюме (RU)", en="Extraction: short summary (EN)")

class LLMKeywordService:
    def __init__(self) -> None:
        pass

    async def generate(self, text: str) -> KeywordTreeSummary:
        await asyncio.sleep(0.08)
        node = KeywordNode(keyword="llm_root", children=[KeywordNode(keyword="llm_child", children=[])])
        return KeywordTreeSummary(ru=node, en=node)

class ExtractionKeywordService:
    def __init__(self) -> None:
        pass

    async def generate(self, text: str) -> KeywordTreeSummary:
        await asyncio.sleep(0.03)
        node = KeywordNode(keyword="extr_root", children=[KeywordNode(keyword="extr_child", children=[])])
        return KeywordTreeSummary(ru=node, en=node)

# ----------------------------
# Coordinator service
# ----------------------------
class SummaryGenerationService:
    def __init__(
        self,
        llm_text_svc: LLMTextSummaryService,
        llm_keyword_svc: LLMKeywordService,
        extraction_text_svc: ExtractionTextSummaryService,
        extraction_keyword_svc: ExtractionKeywordService,
    ) -> None:
        self.llm_text_svc = llm_text_svc
        self.llm_keyword_svc = llm_keyword_svc
        self.extraction_text_svc = extraction_text_svc
        self.extraction_keyword_svc = extraction_keyword_svc

    async def generate_full_summary(self, text: str) -> SummaryResult:
        llm_text, llm_kw, extr_text, extr_kw = await asyncio.gather(
            self.llm_text_svc.generate(text),
            self.llm_keyword_svc.generate(text),
            self.extraction_text_svc.generate(text),
            self.extraction_keyword_svc.generate(text),
        )
        return SummaryResult(
            llm_text_summary=llm_text,
            llm_keyword_summary=llm_kw,
            extraction_text_summary=extr_text,
            extraction_keyword_summary=extr_kw,
        )

# ----------------------------
# File uploader class (new)
# ----------------------------
class FileUploader:
    """Класс для сохранения загруженного файла и извлечения текста (PDF и DOCX).

    Использует синхронные библиотеки для разбора файлов, поэтому heavy work выполняется
    в отдельном потоке через asyncio.to_thread().
    """
    def __init__(self, uploads_dir: Path):
        self.uploads_dir = uploads_dir
        self.uploads_dir.mkdir(parents=True, exist_ok=True)

    async def save_upload(self, upload: UploadFile) -> Path:
        # безопасное имя файла — используем оригинальное имя, но на случай конфликтов добавляем timestamp
        filename = Path(upload.filename).name
        target = self.uploads_dir / f"{int(datetime.utcnow().timestamp())}_{filename}"
        # write in thread
        await asyncio.to_thread(self._write_file, upload, target)
        return target

    def _write_file(self, upload: UploadFile, target_path: Path) -> None:
        # UploadFile.file is a SpooledTemporaryFile-like object; we should seek and copy
        upload.file.seek(0)
        with open(target_path, "wb") as out_f:
            shutil.copyfileobj(upload.file, out_f)
        # After copying, close the upload file to free resources
        try:
            upload.file.close()
        except Exception:
            pass

    async def extract_text(self, file_path: Path) -> str:
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            return await asyncio.to_thread(self._extract_text_pdf, str(file_path))
        elif suffix == ".docx":
            return await asyncio.to_thread(self._extract_text_docx, str(file_path))
        else:
            raise ValueError("Unsupported file type: only PDF and DOCX are allowed")

    def _extract_text_docx(self, path: str) -> str:
        # use python-docx
        try:
            from docx import Document
        except Exception as e:
            raise RuntimeError("python-docx not installed. Install with: pip install python-docx") from e
        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text]
        return "\n\n".join(paragraphs)

    def _extract_text_pdf(self, path: str) -> str:
        # use pdfminer.six extract_text
        try:
            from pdfminer.high_level import extract_text
        except Exception as e:
            raise RuntimeError("pdfminer.six not installed. Install with: pip install pdfminer.six") from e
        text = extract_text(path)
        return text or ""

# ----------------------------
# Document Service (orchestration)
# ----------------------------
class DocumentService:
    def __init__(self, repo: TextRepositoryAsync, summary_service: SummaryGenerationService) -> None:
        self.repo = repo
        self.summary_service = summary_service

    async def create_document(self, file_name: str, text: str, name: Optional[str] = None) -> int:
        summary = await self.summary_service.generate_full_summary(text)
        return await self.repo.add_document(text, summary, file_name, name)

    async def get_document(self, doc_id: int) -> TextDocumentDTO:
        doc = await self.repo.get_document(doc_id)
        if doc is None:
            raise ValueError(f"Документ с id={doc_id} не найден.")
        return doc

    async def list_documents_info(self) -> List[DocumentInfoDTO]:
        return await self.repo.list_document_info()

    async def delete_document(self, doc_id: int) -> bool:
        return await self.repo.delete_document(doc_id)

# ----------------------------
# FastAPI app + DI
# ----------------------------
app = FastAPI(title="Async Document Summary (single-file with Jinja2)")

# mount static and templates
app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")
templates = Jinja2Templates(directory=str(TEMPLATES_DIR))

@app.on_event("startup")
async def startup_event():
    # create repo and services, and init db
    db_url = os.environ.get("DB_URL", "sqlite+aiosqlite:///texts_async.db")
    repo = TextRepositoryAsync(db_url=db_url)
    await repo.init_models()

    # services
    llm_text_svc = LLMTextSummaryService()
    llm_kw_svc = LLMKeywordService()
    extraction_text_svc = ExtractionTextSummaryService()
    extraction_kw_svc = ExtractionKeywordService()
    summary_service = SummaryGenerationService(
        llm_text_svc=llm_text_svc,
        llm_keyword_svc=llm_kw_svc,
        extraction_text_svc=extraction_text_svc,
        extraction_keyword_svc=extraction_kw_svc,
    )
    document_service = DocumentService(repo=repo, summary_service=summary_service)

    # file uploader
    uploader = FileUploader(UPLOADS_DIR)

    app.state.repo = repo
    app.state.document_service = document_service
    app.state.uploader = uploader
    # ensure static and templates exist (in case deployed to empty container)
    ensure_static_and_templates()

@app.on_event("shutdown")
async def shutdown_event():
    repo = getattr(app.state, "repo", None)
    if repo is not None:
        try:
            await repo.engine.dispose()
        except Exception:
            pass

# dependency
def get_document_service(request: Request) -> DocumentService:
    svc = getattr(request.app.state, "document_service", None)
    if svc is None:
        raise HTTPException(status_code=500, detail="DocumentService not initialized")
    return svc

def get_uploader(request: Request) -> FileUploader:
    up = getattr(request.app.state, "uploader", None)
    if up is None:
        raise HTTPException(status_code=500, detail="FileUploader not initialized")
    return up

# ----------------------------
# Web endpoints: HTML (Jinja2) + upload handler
# ----------------------------
@app.get("/", response_class=HTMLResponse)
async def index(request: Request, service: DocumentService = Depends(get_document_service)):
    docs = await service.list_documents_info()
    # format created_at for display
    for d in docs:
        if d.created_at:
            d.created_at = d.created_at.strftime("%Y-%m-%d %H:%M:%S")
    return templates.TemplateResponse("index.html", {"request": request, "documents": docs})

@app.post("/documents/upload")
async def upload_document(
    request: Request,
    file: UploadFile = File(...),
    name: Optional[str] = Form(None),
    service: DocumentService = Depends(get_document_service),
    uploader: FileUploader = Depends(get_uploader),
):
    # save file
    try:
        saved_path = await uploader.save_upload(file)
    except Exception as e:
        return templates.TemplateResponse("error.html", {"request": request, "message": f"Ошибка сохранения файла: {e}"})

    # extract text
    try:
        text = await uploader.extract_text(saved_path)
    except Exception as e:
        return templates.TemplateResponse("error.html", {"request": request, "message": f"Ошибка извлечения текста: {e}"})

    if not text or len(text.strip()) == 0:
        # still allow creating maybe empty text but warn
        text = ""

    # use original filename as file_name (without timestamp prefix)
    orig_filename = "_".join(saved_path.name.split("_")[1:]) if "_" in saved_path.name else saved_path.name

    try:
        doc_id = await service.create_document(file_name=orig_filename, text=text, name=name)
    except Exception as e:
        return templates.TemplateResponse("error.html", {"request": request, "message": f"Ошибка при создании записи: {e}"})

    return RedirectResponse(url=f"/documents/{doc_id}", status_code=status.HTTP_303_SEE_OTHER)

@app.get("/documents/{doc_id}", response_class=HTMLResponse)
async def view_document(doc_id: int, request: Request, service: DocumentService = Depends(get_document_service)):
    try:
        doc = await service.get_document(doc_id)
    except ValueError as e:
        return templates.TemplateResponse("error.html", {"request": request, "message": str(e)})
    # format created_at
    if doc is not None and getattr(doc, "created_at", None):
        try:
            doc.created_at = doc.created_at.strftime("%Y-%m-%d %H:%M:%S")
        except Exception:
            pass
    return templates.TemplateResponse("document.html", {"request": request, "doc": doc})

@app.post("/documents/{doc_id}/delete")
async def delete_document_form(doc_id: int, request: Request, service: DocumentService = Depends(get_document_service)):
    deleted = await service.delete_document(doc_id)
    if not deleted:
        return templates.TemplateResponse("error.html", {"request": request, "message": f"Документ id={doc_id} не найден."})
    return RedirectResponse(url="/", status_code=status.HTTP_303_SEE_OTHER)

# Optional JSON API endpoints
@app.get("/api/documents/", response_model=List[DocumentInfoDTO])
async def api_list_documents(service: DocumentService = Depends(get_document_service)):
    return await service.list_documents_info()

@app.get("/api/documents/{doc_id}", response_model=TextDocumentDTO)
async def api_get_document(doc_id: int, service: DocumentService = Depends(get_document_service)):
    try:
        return await service.get_document(doc_id)
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))

# If run directly
if __name__ == "__main__":
    print("Run with: uvicorn app_full_fixed:app --reload")
